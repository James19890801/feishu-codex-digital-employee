import json
import sys
from html import escape
from pathlib import Path
from docx import Document
from docx.shared import Pt


def write_pdf(target, payload):
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        'ChineseNormal', parent=styles['Normal'], fontName='STSong-Light',
        fontSize=10.5, leading=17, spaceAfter=7,
    )
    title = ParagraphStyle(
        'ChineseTitle', parent=normal, fontSize=20, leading=28,
        alignment=TA_CENTER, spaceAfter=16,
    )
    headings = {
        1: ParagraphStyle('ChineseH1', parent=normal, fontSize=16, leading=24, spaceBefore=12, spaceAfter=8),
        2: ParagraphStyle('ChineseH2', parent=normal, fontSize=14, leading=21, spaceBefore=10, spaceAfter=7),
        3: ParagraphStyle('ChineseH3', parent=normal, fontSize=12, leading=19, spaceBefore=8, spaceAfter=6),
    }
    story = [Paragraph(escape(payload.get('title') or 'AI 数字分身交付物'), title)]
    for raw in payload.get('content', '').splitlines():
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 3 * mm))
        elif line.startswith('### '):
            story.append(Paragraph(escape(line[4:]), headings[3]))
        elif line.startswith('## '):
            story.append(Paragraph(escape(line[3:]), headings[2]))
        elif line.startswith('# '):
            story.append(Paragraph(escape(line[2:]), headings[1]))
        elif line.startswith(('- ', '* ')):
            story.append(Paragraph(f'• {escape(line[2:])}', normal))
        else:
            story.append(Paragraph(escape(line), normal))
    document = SimpleDocTemplate(
        str(target), pagesize=A4,
        rightMargin=18 * mm, leftMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=payload.get('title') or 'AI 数字分身交付物',
    )
    document.build(story)


def main():
    payload = json.loads(sys.stdin.read())
    target = Path(payload['path'])
    target.parent.mkdir(parents=True, exist_ok=True)
    if str(payload.get('format') or '').lower() == 'pdf' or target.suffix.lower() == '.pdf':
        write_pdf(target, payload)
        print(target)
        return
    doc = Document()
    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(11)
    doc.add_heading(payload.get('title') or 'AI 数字分身交付物', 0)
    for raw in payload.get('content', '').splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(('- ', '* ')):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif len(line) > 2 and line[0].isdigit() and line[1] in '.、':
            doc.add_paragraph(line[2:].strip(), style='List Number')
        else:
            doc.add_paragraph(line)
    doc.save(target)
    print(target)


if __name__ == '__main__':
    main()
